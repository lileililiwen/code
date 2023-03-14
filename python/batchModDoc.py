import os
from docx import Document
from docx.shared import Cm

# 待处理的文件夹路径
folder_path = 'path/to/folder'

# 待添加的图片路径
image_path = 'path/to/image.jpg'

# 待添加的文字
text = 'Hello, World!'

# 遍历文件夹下所有Word文件
for filename in os.listdir(folder_path):
    if filename.endswith('.docx'):
        file_path = os.path.join(folder_path, filename)

        # 打开Word文件
        document = Document(file_path)

        # 获取第2页
        page = document.sections[0].add_page_break()
        page = document.add_page_break()

        # 在页面末尾添加图片
        paragraph = page.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_path, width=Cm(10))

        # 在图片下方添加文字
        paragraph = page.add_paragraph(text)

        # 保存Word文件
        document.save(file_path)
#上面的代码会遍历指定文件夹下所有扩展名为.docx的Word文件，并将图片文件image.jpg添加到第2页末尾，下方添加文字Hello, World!，最后保存修改后的Word文件。
#注意：上面的代码只是一个示例，实际应用中可能需要对不同的Word文件进行不同的操作，或者对文件读写、图片添加、文字添加等操作进行异常处理。        