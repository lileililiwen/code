# 安装python-docx库：pip install python-docx

import docx

# 打开Word文档
doc = docx.Document('my_doc.docx')

# 获取文档中的所有段落
paragraphs = doc.paragraphs

# 遍历所有段落并打印
for paragraph in paragraphs:
    print(paragraph.text)

# 获取文档中的所有表格
tables = doc.tables

# 遍历所有表格并打印
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)

# 创建新的Word文档
new_doc = docx.Document()

# 添加新的段落到文档中
new_paragraph = new_doc.add_paragraph('This is a new paragraph.')

# 添加新的表格到文档中
new_table = new_doc.add_table(rows=3, cols=3)
for i in range(3):
    for j in range(3):
        new_table.cell(i, j).text = f'Row {i}, Column {j}'

# 保存文档
new_doc.save('new_doc.docx')